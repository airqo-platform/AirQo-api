from google.cloud import storage
from google.cloud import bigquery
import pandas as pd
import numpy as np
import seaborn as sns

sns.set_style("whitegrid")
import matplotlib.pyplot as plt

plt.rcParams["figure.figsize"] = [14, 6]
import matplotlib.dates as mdates
from docx import Document
from docx.shared import Inches
import io
import ast
import gcsfs
from config import environment, configuration
import os
import joblib
import datetime as dt
from datetime import datetime, timedelta
from utils import (
    upload_trained_model_to_gcs,
    date_to_str,
    str_to_date_2,
    upload_csv_file_to_gcs,
    previous_months_range,
)
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import warnings

warnings.filterwarnings("ignore")


BIGQUERY_SITES = configuration.BIGQUERY_SITES
BIGQUERY_SITES_METADATA = configuration.BIGQUERY_SITES_METADATA
BIGQUERY_EVENTS = configuration.BIGQUERY_EVENTS
POLLUTANT_BIGQUERY_MAPPER = {
    "pm2_5": ["pm2_5_calibrated_value", "pm2_5_raw_value"],
    "pm10": ["pm10_calibrated_value", "pm10_raw_value"],
}

expected_days = int(configuration.EXPECTED_DAYS)
expected_hourly_records = expected_days * 24
cutoff_count = 0.75 * expected_hourly_records
revised_cutoff_count = 0.60 * expected_hourly_records


def get_city_data(data_):
    data_.rename(columns={"datetime": "timestamp"}, inplace=True)
    data_["timestamp"] = pd.to_datetime(data_["timestamp"])
    data_["timestamp"] = pd.to_datetime(
        data_["timestamp"].dt.strftime("%Y-%m-%d %H:%M:%S")
    )
    data_["timestamp"] = data_["timestamp"].apply(lambda x: x.replace(tzinfo=None))
    data_ = data_[data_["pm2_5_calibrated_value"] <= 500]
    data_ = data_[data_["pm2_5_calibrated_value"] >= 1]
    data_["month"] = data_["timestamp"].dt.month
    data_["year"] = data_["timestamp"].dt.year
    data_["hour"] = data_["timestamp"].dt.hour
    data_["day_of_week"] = data_["timestamp"].dt.strftime("%A")
    data_["Month"] = data_["timestamp"].dt.strftime("%b")
    data = data_[
        [
            "timestamp",
            "site",
            "pm2_5_calibrated_value",
            "pm2_5_raw_value",
            "pm10_calibrated_value",
            "pm10_raw_value",
            "city",
            "month",
            "year",
            "Month",
        ]
    ]
    return data


def get_yearly(filtered_df, col_value="pm2_5_calibrated_value"):
    yearly = (
        filtered_df[[col_value, "year"]].groupby(filtered_df["year"]).mean().round(2)
    )
    yearly["Year"] = yearly.index
    yearly["Year"] = yearly["Year"].astype(str)
    return yearly


def plot_yearly(yearly, city, col_value="pm2_5_calibrated_value"):
    f, a = plt.subplots(figsize=(8, 5))
    yn = sns.barplot(x=yearly.Year, y=col_value, data=yearly, axes=a)
    y_axis_label = "PM$_{2.5}$ Concentration(µg/m\N{SUPERSCRIPT THREE})"
    plt.ylabel(y_axis_label)
    plt.title("Average Yearly PM2.5 Concentration in " + city)
    mem_image_file = io.BytesIO()
    ww = plt.savefig(mem_image_file, format="png")
    plt.close()
    return f, mem_image_file


def get_monthly(filtered_df, col_value="pm2_5_calibrated_value"):
    monthly = (
        filtered_df[[col_value, "Month"]].groupby(filtered_df["Month"]).mean().round(2)
    )
    ordered_months = [
        "Jan",
        "Feb",
        "Mar",
        "Apr",
        "May",
        "Jun",
        "Jul",
        "Aug",
        "Sep",
        "Oct",
        "Nov",
        "Dec",
    ]
    monthly = monthly.reindex(ordered_months)
    monthly_ = monthly.reset_index()
    return monthly_


def plot_monthly(monthly, city, col_value="pm2_5_calibrated_value"):
    figm, axm = plt.subplots(figsize=(10, 6))
    mn = sns.barplot(x=monthly.index, y=col_value, data=monthly, axes=axm)
    y_axis_label = "PM$_{2.5}$ Concentration(µg/m\N{SUPERSCRIPT THREE})"
    plt.ylabel(y_axis_label)
    plt.title("Average Monthly PM2.5 Concentration in " + city)
    mem_image_file = io.BytesIO()
    ww = plt.savefig(mem_image_file, format="png")
    plt.close()
    return figm, mem_image_file


def sort_sites_based_on(
    df, index="site", col_values="pm2_5_calibrated_value", order=False
):
    df_p = df.pivot_table(index=index, values=col_values)
    sorted_df = df_p.reset_index().sort_values(by=col_values, ascending=order)
    sorted_df = sorted_df.reset_index().sort_values(by=col_values, ascending=False)
    sorted_df.drop("index", axis=1, inplace=True)
    return sorted_df.round(2)


def visualise_data_(df, device_id, label_value, column_value, ax):
    df.sort_values("timestamp", ascending=True)
    df2 = df.resample("M").mean().round(2).dropna(axis=0)
    date_form = mdates.DateFormatter("%Y-%m")
    monthly_interval = mdates.MonthLocator(interval=4)

    ax.plot(df2.index, df2[column_value], label=label_value)
    ax.set_title(device_id)
    ax.set(xlabel="Date", ylabel=label_value)
    ax.xaxis.set_major_locator(monthly_interval)
    ax.xaxis.set_major_formatter(date_form)
    ax.tick_params(axis="x", labelrotation=60)
    plt.legend()


def plot_bar_graph(data, city, x_label_rotation=90):
    ax = sns.barplot(
        x="site", y="pm2_5_calibrated_value", data=data, ci=None, palette="Blues"
    )
    plt.title("Average PM$_{2.5}$ Concentration in " + city, fontsize=28)
    plt.xlabel("Location ")
    plt.ylabel("PM$_{2.5}$ $\mu$ g/m$^3$")
    plt.xticks(rotation=x_label_rotation)
    plt.show()


def reprocess_to_monthly_df(df):
    df2 = df.resample("M").mean().round(2).dropna(axis=0).reset_index()
    return df2


def reprocess_to_daily_df(df):
    df2 = df.resample("D").mean().round(2).dropna(axis=0).reset_index()
    return df2


def reprocess_to_annual_df(df):
    df_ = df.resample("D").mean().round(2).dropna(axis=0)
    df2 = df_.resample("A").mean().round(2).dropna(axis=0).reset_index()
    return df2


def convert_AQI_category(row):
    if row["pm2_5_calibrated_value"] <= 12.0:
        return "Good"
    elif row["pm2_5_calibrated_value"] <= 35.4:
        return "Moderate"
    elif row["pm2_5_calibrated_value"] <= 55.4:
        return "Unhealthy for sensitive group"
    elif row["pm2_5_calibrated_value"] <= 150.4:
        return "Unhealthy"
    elif row["pm2_5_calibrated_value"] <= 250.4:
        return "Very Unhealthy"
    else:
        return "Hazardous"


def compute_percentage(expected_records, actual_records):
    print(expected_records)
    print(actual_records)
    if actual_records != 0:
        return round((actual_records / expected_records) * 100, 2)
    else:
        return 0


def aggregate_data(
    data_unique,
    col_value="pm2_5_calibrated_value",
    second_col_value="pm2_5_raw_value",
    third_col_value="pm10_calibrated_value",
    fourth_col_value="pm10_raw_value",
    total_expected_records_in_the_timeperiod=2160,
):
    device_records = []
    results_dictionary = []
    unique_sites = data_unique.site.unique()
    for device in unique_sites:
        kd = data_unique[data_unique["site"] == device]
        res = (
            kd[[col_value, second_col_value, third_col_value, fourth_col_value]]
            .describe()
            .round(2)
        )
        device_records.append(kd)
        overall_percentage_of_available_calibrated_pm2_5 = compute_percentage(
            total_expected_records_in_the_timeperiod,
            int(res[[col_value]].loc["count"].values[0]),
        )

        overall_percentage_of_available_raw_pm2_5 = compute_percentage(
            total_expected_records_in_the_timeperiod,
            int(res[[second_col_value]].loc["count"].values[0]),
        )

        overall_percentage_of_missing_calibrated_pm2_5 = (
            100 - overall_percentage_of_available_calibrated_pm2_5
        )
        overall_percentage_of_missing_raw_pm2_5 = (
            100 - overall_percentage_of_available_raw_pm2_5
        )

        overall_percentage_of_available_calibrated_pm10 = compute_percentage(
            total_expected_records_in_the_timeperiod,
            int(res[[third_col_value]].loc["count"].values[0]),
        )

        overall_percentage_of_available_raw_pm10 = compute_percentage(
            total_expected_records_in_the_timeperiod,
            int(res[[fourth_col_value]].loc["count"].values[0]),
        )

        overall_percentage_of_missing_calibrated_pm10 = (
            100 - overall_percentage_of_available_calibrated_pm10
        )
        overall_percentage_of_missing_raw_pm10 = (
            100 - overall_percentage_of_available_raw_pm10
        )

        dic = {
            "site": device,
            "calibrated_pm2_5_count": int(res[[col_value]].loc["count"].values[0]),
            "avg_calibrated_pm2_5": res[[col_value]].loc["mean"].values[0],
            "min_calibrated_pm2_5": res[[col_value]].loc["min"].values[0],
            "max_calibrated_pm2_5": res[[col_value]].loc["max"].values[0],
            "std_calibrated_pm2_5": res[[col_value]].loc["std"].values[0],
            "raw_pm2_5_count": int(res[[second_col_value]].loc["count"].values[0]),
            "avg_raw_pm2_5": res[[second_col_value]].loc["mean"].values[0],
            "min_raw_pm2_5": res[[second_col_value]].loc["min"].values[0],
            "max_raw_pm2_5": res[[second_col_value]].loc["max"].values[0],
            "std_raw_pm2_5": res[[second_col_value]].loc["std"].values[0],
            "calibrated_pm10_count": int(res[[third_col_value]].loc["count"].values[0]),
            "avg_calibrated_pm10": res[[third_col_value]].loc["mean"].values[0],
            "min_calibrated_pm10": res[[third_col_value]].loc["min"].values[0],
            "max_calibrated_pm10": res[[third_col_value]].loc["max"].values[0],
            "std_calibrated_pm10": res[[third_col_value]].loc["std"].values[0],
            "raw_pm10_count": int(res[[fourth_col_value]].loc["count"].values[0]),
            "avg_raw_pm10": res[[fourth_col_value]].loc["mean"].values[0],
            "min_raw_pm10": res[[fourth_col_value]].loc["min"].values[0],
            "max_raw_pm10": res[[fourth_col_value]].loc["max"].values[0],
            "std_raw_pm10": res[[fourth_col_value]].loc["std"].values[0],
            "percentage_of_calibrated_pm2_5": compute_percentage(
                int(res[[second_col_value]].loc["count"].values[0]),
                int(res[[col_value]].loc["count"].values[0]),
            ),
            "percentage_of_calibrated_pm10": compute_percentage(
                int(res[[fourth_col_value]].loc["count"].values[0]),
                int(res[[third_col_value]].loc["count"].values[0]),
            ),
            "overall_percentage_of_available_calibrated_pm2_5": overall_percentage_of_available_calibrated_pm2_5,
            "overall_percentage_of_missing_calibrated_pm2_5": overall_percentage_of_missing_calibrated_pm2_5,
            "overall_percentage_of_available_calibrated_pm10": overall_percentage_of_available_calibrated_pm10,
            "overall_percentage_of_missing_calibrated_pm10": overall_percentage_of_missing_calibrated_pm10,
            "overall_percentage_of_available_raw_pm2_5": overall_percentage_of_available_raw_pm2_5,
            "overall_percentage_of_missing_raw_pm2_5": overall_percentage_of_missing_raw_pm2_5,
            "overall_percentage_of_available_raw_pm10": overall_percentage_of_available_raw_pm10,
            "overall_percentage_of_missing_raw_pm10": overall_percentage_of_missing_raw_pm10,
        }
        results_dictionary.append(dic)

    final_df = pd.DataFrame(results_dictionary)
    final_df.sort_values(
        by=["percentage_of_calibrated_pm2_5"], ascending=False, inplace=True
    )
    return final_df


def get_sites_to_omit_due_to_fewer_observations(final_df, count=2000):
    omitted_sites_df = final_df[final_df.raw_pm2_5_count <= count]
    omitted_sites_df.reset_index(drop=True)
    omitted_sites = omitted_sites_df["site"].unique()
    print(len(omitted_sites))
    return omitted_sites


def filter_sites_to_use(data_unique, omitted_sites):
    filtered_df = data_unique[~data_unique["site"].isin(omitted_sites)]
    return filtered_df


def get_site_summarised_stats(data, expected_hourly_records):
    site_summarised_data = aggregate_data(
        data,
        "pm2_5_calibrated_value",
        "pm2_5_raw_value",
        "pm10_calibrated_value",
        "pm10_raw_value",
        expected_hourly_records,
    )

    site_descriptive_stats_calibrated_pm2_5 = site_summarised_data[
        [
            "site",
            "avg_calibrated_pm2_5",
            "min_calibrated_pm2_5",
            "max_calibrated_pm2_5",
            "std_calibrated_pm2_5",
        ]
    ].round(2)

    site_descriptive_stats_calibrated_pm2_5.rename(
        columns={
            "site": "site name",
            "avg_calibrated_pm2_5": "Mean",
            "min_calibrated_pm2_5": "Min",
            "max_calibrated_pm2_5": "Max",
            "std_calibrated_pm2_5": "STD",
        },
        inplace=True,
    )

    site_descriptive_stats_calibrated_pm10 = site_summarised_data[
        [
            "site",
            "avg_calibrated_pm10",
            "min_calibrated_pm10",
            "max_calibrated_pm10",
            "std_calibrated_pm10",
        ]
    ].round(2)

    site_descriptive_stats_calibrated_pm10.rename(
        columns={
            "site": "site name",
            "avg_calibrated_pm10": "Mean",
            "min_calibrated_pm10": "Min",
            "max_calibrated_pm10": "Max",
            "std_calibrated_pm10": "STD",
        },
        inplace=True,
    )

    site_descriptive_stats_raw_pm2_5 = site_summarised_data[
        ["site", "avg_raw_pm2_5", "min_raw_pm2_5", "max_raw_pm2_5", "std_raw_pm2_5"]
    ].round(2)

    site_descriptive_stats_raw_pm2_5.rename(
        columns={
            "site": "site name",
            "avg_raw_pm2_5": "Mean",
            "min_raw_pm2_5": "Min",
            "max_raw_pm2_5": "Max",
            "std_raw_pm2_5": "STD",
        },
        inplace=True,
    )

    site_descriptive_stats_raw_pm10 = site_summarised_data[
        ["site", "avg_raw_pm10", "min_raw_pm10", "max_raw_pm10", "std_raw_pm10"]
    ].round(2)

    site_descriptive_stats_raw_pm10.rename(
        columns={
            "site": "site name",
            "avg_raw_pm10": "Mean",
            "min_raw_pm10": "Min",
            "max_raw_pm10": "Max",
            "std_raw_pm10": "STD",
        },
        inplace=True,
    )

    sorted_site = sort_sites_based_on(
        data, index="site", col_values="pm2_5_calibrated_value", order=False
    )

    return (
        site_summarised_data,
        site_descriptive_stats_calibrated_pm2_5,
        site_descriptive_stats_calibrated_pm10,
        site_descriptive_stats_calibrated_pm10,
        site_descriptive_stats_raw_pm2_5,
        site_descriptive_stats_raw_pm10,
        sorted_site,
    )


def generate_site_stats_dictionaries_for_saving_to_file(
    site_calibrated_data_summary,
    overall_site_calibrated,
    overall_site_raw,
    site_descriptive_stats_calibrated_pm2_5,
    site_descriptive_stats_calibrated_pm10,
    site_descriptive_stats_raw_pm2_5,
    site_descriptive_stats_raw_pm10,
    sites_to_omit,
    sorted_sites,
):
    site_overall_raw_data_completeness_ = {
        "title": "Table showing overall data completeness based on raw measurements",
        "data": overall_site_raw,
    }

    site_overall_calibrated_data_completeness_ = {
        "title": "Table showing overall data completeness based on calibrated measurements",
        "data": overall_site_calibrated,
    }

    site_calibrated_data_summary_ = {
        "title": "Table showing percentage of calibrated data in relation to raw measurements",
        "data": site_calibrated_data_summary,
    }

    site_calibrated_pm2_5_descriptive_stat_ = {
        "title": "Table showing descriptive statistics for calibrated PM2.5 data",
        "data": site_descriptive_stats_calibrated_pm2_5,
    }

    site_calibrated_pm10_descriptive_stat_ = {
        "title": "Table showing descriptive statistics for calibrated PM10 data",
        "data": site_descriptive_stats_calibrated_pm10,
    }

    site_raw_pm2_5_descriptive_stat_ = {
        "title": "Table showing descriptive statistics for raw PM2.5 data",
        "data": site_descriptive_stats_raw_pm2_5,
    }

    site_raw_pm10_descriptive_stat_ = {
        "title": "Table showing descriptive statistics for raw PM10 data",
        "data": site_descriptive_stats_raw_pm10,
    }

    sites_to_omit_heading_ = {
        "content_type": "heading",
        "text": "Number of sites to omit i.e. not meeting the required completennes threshold",
    }
    sites_to_omit_text_ = {
        "content_type": "text_content",
        "text": str(len(sites_to_omit))
        + " sites do not meet the required completennes threshold i.e;",
    }
    sites_to_omit_ = {
        "content_type": "list",
        "text": "Sites to omit i.e. not meeting the required completennes threshold",
        "list_items": sites_to_omit,
    }

    srt_ranked = {
        "title": "Table showing monitoring sites average air quality ranked in descending order",
        "data": sorted_sites,
    }

    return (
        site_overall_raw_data_completeness_,
        site_overall_calibrated_data_completeness_,
        site_calibrated_data_summary_,
        site_calibrated_pm2_5_descriptive_stat_,
        site_calibrated_pm10_descriptive_stat_,
        site_raw_pm2_5_descriptive_stat_,
        site_raw_pm10_descriptive_stat_,
        srt_ranked,
        sites_to_omit_heading_,
        sites_to_omit_text_,
        sites_to_omit_,
    )


def get_site_data_completenness_stats(site_summarised_data):
    site_calibrated_data_summary = site_summarised_data[
        [
            "site",
            "calibrated_pm2_5_count",
            "raw_pm2_5_count",
            "calibrated_pm10_count",
            "raw_pm10_count",
            "percentage_of_calibrated_pm2_5",
            "percentage_of_calibrated_pm10",
        ]
    ]

    site_calibrated_data_summary.rename(
        columns={
            "site": "site name",
            "calibrated_pm2_5_count": "Calibrated PM2.5 Count",
            "raw_pm2_5_count": "Raw PM2.5 Count",
            "calibrated_pm10_count": "Calibrated PM10 Count",
            "raw_pm10_count": "Raw PM10 Count",
            "percentage_of_calibrated_pm2_5": "Calibrated PM2.5(%)",
            "percentage_of_calibrated_pm10": "Calibrated PM10(%)",
        },
        inplace=True,
    )

    overall_site_calibrated = site_summarised_data[
        [
            "site",
            "overall_percentage_of_available_calibrated_pm2_5",
            "overall_percentage_of_missing_calibrated_pm2_5",
            "overall_percentage_of_available_calibrated_pm10",
            "overall_percentage_of_missing_calibrated_pm10",
        ]
    ]

    overall_site_calibrated.rename(
        columns={
            "site": "site name",
            "overall_percentage_of_available_calibrated_pm2_5": "Completeness(%) PM2.5",
            "overall_percentage_of_missing_calibrated_pm2_5": "Missing(%) PM2.5",
            "overall_percentage_of_available_calibrated_pm10": "Completeness(%) PM10",
            "overall_percentage_of_missing_calibrated_pm10": "Missing(%) PM10",
        },
        inplace=True,
    )

    overall_site_raw = site_summarised_data[
        [
            "site",
            "overall_percentage_of_available_raw_pm2_5",
            "overall_percentage_of_missing_raw_pm2_5",
            "overall_percentage_of_available_raw_pm10",
            "overall_percentage_of_missing_raw_pm10",
        ]
    ]

    overall_site_raw.rename(
        columns={
            "site": "site name",
            "overall_percentage_of_available_raw_pm2_5": "Completeness(%) PM2.5",
            "overall_percentage_of_missing_raw_pm2_5": "Missing(%) PM2.5",
            "overall_percentage_of_available_raw_pm10": "Completeness(%) PM10",
            "overall_percentage_of_missing_raw_pm10": "Missing(%) PM10",
        },
        inplace=True,
    )

    sites_to_omit = get_sites_to_omit_due_to_fewer_observations(
        site_summarised_data, cutoff_count
    )

    return (
        site_calibrated_data_summary,
        overall_site_calibrated,
        overall_site_raw,
        sites_to_omit,
    )


def save_monthly_and_annual_analysis_for_location(report_path, location_data):
    for location_df in location_data:
        location_name = location_df["site"].iat[0]
        location_name_heading_ = {
            "content_type": "heading",
            "text": location_name + " overview",
        }
        location_monthly_data = get_monthly(location_df)
        location_monthly_data_ = {
            "content_type": "table",
            "title": "Table showing "
            + location_name
            + " monthly averages for calibrated PM2.5 data",
            "data": location_monthly_data,
        }

        monthly_plot_fig, monthly_plot_mem_fig = plot_monthly(
            location_monthly_data, location_name
        )

        location_monthly_plot_ = {
            "content_type": "image",
            "image_label": "Figure showing "
            + location_name
            + " monthly averages for calibrated PM2.5 data",
            "image": monthly_plot_mem_fig,
        }

        location_yearly_data = get_yearly(location_df)
        location_yearly_data_ = {
            "content_type": "table",
            "title": "Table showing "
            + location_name
            + " 24 hour annual averages for calibrated PM2.5 data",
            "data": location_yearly_data,
        }

        yearly_plot_fig, yearly_plot_mem_fig = plot_yearly(
            location_yearly_data, location_name
        )

        location_yearly_plot_ = {
            "content_type": "image",
            "image_label": "Figure showing "
            + location_name
            + " 24 hour annual averages for calibrated PM2.5 data",
            "image": yearly_plot_mem_fig,
        }

        write_city_analysis_to_word_file(
            report_path,
            location_name_heading=location_name_heading_,
            location_monthly_data=location_monthly_data_,
            location_monthly_plot_=location_monthly_plot_,
            location_yearly_data=location_yearly_data_,
            location_yearly_plot_=location_yearly_plot_,
        )


def write_report_template(file_path=None):
    if file_path == None:
        document = Document()
        file_path = "demo.docx"
    else:
        document = Document(file_path)

    document.add_heading("Quarterly Air Quality Report Analysis", 0)

    p = document.add_paragraph("Analysis to be used for generating quarterly report ")
    p.add_run("bold").bold = True
    p.add_run(" and some ")
    p.add_run("italic.").italic = True

    document.add_heading("Data Completeness Checks For all the cities", level=1)
    document.add_paragraph("How completeness was computed", style="Intense Quote")

    document.add_paragraph("Expected days in the quarter is 90", style="List Bullet")
    document.add_paragraph(
        "Expected hourly records in the two years is expected days*24",
        style="List Bullet",
    )
    document.add_paragraph(
        "Cutoff Count is : 0.75*expected hourly records (75%)", style="List Bullet"
    )

    document.save("demo.docx")


def write_city_analysis_to_word_file(file_path=None, city=None, *args_df, **kwargs):
    if file_path == None:
        document = Document()
        file_path = "demo.docx"
    else:
        document = Document(file_path)

    if city != None:
        document.add_heading(city, level=2)

    for dic in args_df:
        content_df = dic.get("data")
        title = dic.get("title")

        table_ = document.add_table(
            content_df.shape[0] + 1, content_df.shape[1], style="Table Grid"
        )

        # add the header rows.
        for j in range(content_df.shape[-1]):
            table_.cell(0, j).text = content_df.columns[j]

        for i in range(content_df.shape[0]):
            for j in range(content_df.shape[-1]):
                table_.cell(i + 1, j).text = str(content_df.values[i][j])

        document.add_paragraph(title)

    for key, value in kwargs.items():
        content_type = value.get("content_type")

        if content_type == "image":
            image = value.get("image")
            image_label = value.get("image_label")
            document.add_picture(image, width=Inches(6))
            document.add_paragraph(image_label)
        elif content_type == "heading":
            text_content = value.get("text")
            document.add_heading(text_content, level=1)

        elif content_type == "list":
            list_items = value.get("list_items")
            for x in list_items:
                document.add_paragraph(x, style="List Bullet")

        elif content_type == "table":
            content_df = value.get("data")
            table_title = value.get("title")
            table_ = document.add_table(
                content_df.shape[0] + 1, content_df.shape[1], style="Table Grid"
            )
            for j in range(content_df.shape[-1]):
                table_.cell(0, j).text = content_df.columns[j]

            for i in range(content_df.shape[0]):
                for j in range(content_df.shape[-1]):
                    table_.cell(i + 1, j).text = str(content_df.values[i][j])

            document.add_paragraph(table_title)

        else:
            text_content = value.get("text")
            extra_content_value = value.get("data")
            document.add_paragraph(text_content)
            if extra_content_value != "":
                document.add_paragraph(extra_content_value)

    document.save(file_path)


def generate_all_cities_analysis(data, report_path):
    data.city = data.city.fillna("Missing City")
    unique_cities = data["city"].unique()
    for city in unique_cities:
        city_data = data[data["city"] == city]
        kapchorwa_data = get_city_data(city_data)
        if kapchorwa_data.shape[0] > 1:
            (
                kapchorwa_site_summarised_data,
                site_descriptive_stats_calibrated_pm2_5,
                site_descriptive_stats_calibrated_pm10,
                site_descriptive_stats_calibrated_pm10,
                site_descriptive_stats_raw_pm2_5,
                site_descriptive_stats_raw_pm10,
                sorted_site,
            ) = get_site_summarised_stats(kapchorwa_data, expected_hourly_records)
            (
                site_calibrated_data_summary,
                overall_site_calibrated,
                overall_site_raw,
                kapchorwa_sites_to_omit,
            ) = get_site_data_completenness_stats(kapchorwa_site_summarised_data)

            (
                site_overall_raw_data_completeness_,
                site_overall_calibrated_data_completeness_,
                site_calibrated_data_summary_,
                site_calibrated_pm2_5_descriptive_stat_,
                site_calibrated_pm10_descriptive_stat_,
                site_raw_pm2_5_descriptive_stat_,
                site_raw_pm10_descriptive_stat_,
                srt_ranked,
                sites_to_omit_heading_,
                sites_to_omit_text_,
                sites_to_omit_,
            ) = generate_site_stats_dictionaries_for_saving_to_file(
                site_calibrated_data_summary,
                overall_site_calibrated,
                overall_site_raw,
                site_descriptive_stats_calibrated_pm2_5,
                site_descriptive_stats_calibrated_pm10,
                site_descriptive_stats_raw_pm2_5,
                site_descriptive_stats_raw_pm10,
                kapchorwa_sites_to_omit,
                sorted_site,
            )
            write_city_analysis_to_word_file(
                report_path,
                city,
                site_overall_raw_data_completeness_,
                site_overall_calibrated_data_completeness_,
                site_calibrated_data_summary_,
                site_calibrated_pm2_5_descriptive_stat_,
                site_calibrated_pm10_descriptive_stat_,
                site_raw_pm2_5_descriptive_stat_,
                site_raw_pm10_descriptive_stat_,
                srt_ranked,
                sites_to_omit_heading=sites_to_omit_heading_,
                sites_to_omit_size=sites_to_omit_text_,
                sites_to_omit=sites_to_omit_,
            )


def from_bigquery(
    tenant,
    start_date,
    end_date,
    frequency,
    pollutants,
    additional_columns=None,
):
    if additional_columns is None:
        additional_columns = []

    decimal_places = 2

    columns_ = [
        f"{BIGQUERY_EVENTS}.device_id AS device",
        f"{BIGQUERY_SITES}.name AS site",
        "FORMAT_DATETIME('%Y-%m-%d %H:%M:%S', timestamp) AS datetime",
        f"{BIGQUERY_SITES}.approximate_latitude AS latitude",
        f"{BIGQUERY_SITES}.approximate_longitude  AS longitude",
    ]

    columns = [
        f"{BIGQUERY_EVENTS}.device_id AS device",
        f"{BIGQUERY_SITES}.name AS site",
        "FORMAT_DATETIME('%Y-%m-%d %H:%M:%S', timestamp) AS datetime",
        f"{BIGQUERY_SITES}.approximate_latitude AS latitude",
        f"{BIGQUERY_SITES}.approximate_longitude  AS longitude",
        f"{BIGQUERY_SITES}.city AS city",
        f"{BIGQUERY_SITES}.region AS region",
        # f"{BIGQUERY_SITES_METADATA}.sub_county AS subcounty",
        # f"{BIGQUERY_SITES_METADATA}.district AS district",
        # f"{BIGQUERY_SITES_METADATA}.parish AS parish",
        f"{BIGQUERY_SITES}.country AS country",
    ]
    columns.extend(additional_columns)

    for pollutant in pollutants:
        pollutant_mapping = POLLUTANT_BIGQUERY_MAPPER.get(pollutant, [])
        columns.extend(
            [
                f"ROUND({mapping}, {decimal_places}) AS {mapping}"
                for mapping in pollutant_mapping
            ]
        )

    QUERY = (
        f"SELECT {', '.join(map(str, set(columns)))} "
        f"FROM {BIGQUERY_EVENTS} "
        f"JOIN {BIGQUERY_SITES} ON {BIGQUERY_SITES}.id = {BIGQUERY_EVENTS}.site_id "
        f"WHERE {BIGQUERY_EVENTS}.tenant = '{tenant}' "
        f"AND {BIGQUERY_EVENTS}.timestamp >= '{start_date}' "
        f"AND {BIGQUERY_EVENTS}.timestamp <= '{end_date}' "
    )

    job_config = bigquery.QueryJobConfig()
    job_config.use_query_cache = True

    dataframe = bigquery.Client().query(QUERY, job_config).result().to_dataframe()
    dataframe.sort_values(["site", "datetime", "device"], ascending=True, inplace=True)

    return dataframe


def get_quarterly_data(tenant="airqo"):
    start_date = str_to_date_2("2022-07-01 00:00:00")
    end_date = str_to_date_2("2022-09-30 23:59:00")  # previous_months_range(3)
    frequency = "hourly"
    pollutants = ["pm2_5", "pm10"]

    data = from_bigquery(
        tenant=tenant,
        start_date=start_date,
        end_date=end_date,
        frequency=frequency,
        pollutants=pollutants,
    )
    print(data.shape)

    data.to_csv("quarterly_data.csv")

    return data


def send_email():
    mail_content = """Hello,
  This is the quarterly air quality analytics email.
   We have attached the detailed analysis in a word document.

  Thank You
"""
    # The mail addresses and password
    sender_address = configuration.MAIL_SENDER_EMAILADDRESS
    sender_pass = configuration.MAIL_SENDER_PASSWORD
    receiver_addresses = configuration.MAIL_RECEIVER_EMAILADDRESS
    recipients = receiver_addresses.split(",")
    mail_subject = configuration.MAIL_SUBJECT
    # Setup the MIME
    message = MIMEMultipart()
    message["From"] = sender_address
    message["To"] = ", ".join(recipients)
    message["Subject"] = mail_subject
    # The subject line
    # The body and the attachments for the mail
    message.attach(MIMEText(mail_content, "plain"))
    attach_file_name = (
        "demo.docx"  ##TODO: read it from generated report saved in cloudstorage
    )
    attach_file = open(attach_file_name, "rb")  # Open the file as binary mode
    payload = MIMEBase("application", "octate-stream")
    payload.set_payload((attach_file).read())
    encoders.encode_base64(payload)  # encode the attachment
    print(attach_file_name)
    # add payload header with filename
    payload.add_header("Content-Disposition", "attachment", filename=attach_file_name)
    message.attach(payload)
    # Create SMTP session for sending the mail
    session = smtplib.SMTP("smtp.gmail.com", 587)  # use gmail with port
    session.starttls()  # enable security
    session.login(sender_address, sender_pass)  # login with mail_id and password
    text = message.as_string()
    session.sendmail(sender_address, recipients, text)
    session.quit()
    print("Mail Sent")


if __name__ == "__main__":
    quarterly_data = get_quarterly_data()
    print(quarterly_data.shape)
    write_report_template()
    generate_all_cities_analysis(quarterly_data, "demo.docx")
    send_email()
